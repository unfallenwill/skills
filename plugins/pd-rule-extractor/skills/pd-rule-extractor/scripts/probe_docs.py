# /// script
# requires-python = ">=3.10"
# dependencies = [
#   "openpyxl>=3.1",
#   "python-docx>=1.1",
#   "pypdf>=4.0",
# ]
# ///
"""Probe clinical trial source documents: identify roles and extract text chunks.

Subcommands:
  identify <dir>    Scan a directory for docx/xlsx/pdf files and classify each
                    as protocol / dvp / dbdesign / unknown (JSON to stdout).
  extract           Parse one document into chunks + update doc-map.json in the
                    working directory.

The extraction is idempotent per doc_id: re-running with the same doc_id
replaces that document's entry and its chunk files, leaving others intact.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

# Chunk size target: sections longer than this are split at paragraph
# boundaries; sheets longer than this many rows are split into row windows.
MAX_CHUNK_CHARS = 8000
XLSX_ROWS_PER_CHUNK = 200
PDF_PAGES_PER_CHUNK = 5

# Keyword tables for heuristic classification. Both Chinese and English
# spellings are covered because source documents vary by sponsor/CRO.
PROTOCOL_NAME_KW = ["方案", "protocol", "研究方案", "试验方案"]
DVP_NAME_KW = ["dvp", "数据核查", "核查计划", "data validation", "edit check", "数据审核"]
DBDESIGN_NAME_KW = [
    "数据库设计", "dbdesign", "db_design", "database design",
    "crf", "ecrf", "annotat", "建库", "数据点", "数据说明", "sds",
]
PROTOCOL_CONTENT_KW = ["入选标准", "排除标准", "inclusion criteria", "exclusion criteria"]
DVP_CONTENT_KW = ["规则编号", "rule id", "rule no", "核查规则", "rule description", "规则描述"]
DBDESIGN_CONTENT_KW = [
    "crf", "ecrf", "annotat", "数据集", "dataset", "数据点",
    "变量名", "variable", "字段定义", "数据结构", "oid",
]
# Header cells that strongly indicate a DVP rule table (rule id column).
DVP_HEADER_KW = ["规则编号", "rule id", "rule no", "rule_id", "rule#", "核查编号", "编号"]


def eprint(msg: str) -> None:
    print(msg, file=sys.stderr)


def fail(msg: str, code: int = 1) -> None:
    eprint(f"error: {msg}")
    sys.exit(code)


# ---------------------------------------------------------------------------
# identify
# ---------------------------------------------------------------------------

def classify_file(path: Path) -> dict:
    """Return a classification record for one document."""
    name = path.name.lower()
    suffix = path.suffix.lower()
    fmt = suffix.lstrip(".")
    role = "unknown"
    signals: list[str] = []

    def name_hits(kws: list[str]) -> list[str]:
        return [kw for kw in kws if kw in name]

    content_text = ""
    try:
        if suffix == ".docx":
            content_text = _docx_sample_text(path)
        elif suffix == ".xlsx":
            content_text = _xlsx_sample_text(path)
        elif suffix == ".pdf":
            content_text = _pdf_sample_text(path)
    except Exception as exc:  # unreadable file -> unknown, but say why
        signals.append(f"content read failed: {exc}")

    content_lower = content_text.lower()

    proto_score = len(name_hits(PROTOCOL_NAME_KW)) * 2
    dvp_score = len(name_hits(DVP_NAME_KW)) * 2
    dbd_score = len(name_hits(DBDESIGN_NAME_KW)) * 2

    proto_hits = [kw for kw in PROTOCOL_CONTENT_KW if kw in content_lower]
    dvp_hits = [kw for kw in DVP_CONTENT_KW if kw in content_lower]
    dbd_hits = [kw for kw in DBDESIGN_CONTENT_KW if kw in content_lower]
    proto_score += len(proto_hits) * 3
    dvp_score += len(dvp_hits) * 3
    dbd_score += len(dbd_hits) * 3

    # A DVP is almost always an xlsx whose header row contains a rule-id-like
    # column; give that signal decisive weight. Generic xlsx gets a small bump.
    if suffix == ".xlsx":
        dvp_score += 1
        header_line = content_lower.splitlines()[1] if "\n" in content_lower else content_lower
        if any(kw in header_line for kw in DVP_HEADER_KW):
            dvp_score += 4
            signals.append("header row contains a rule-id column")

    scores = {"protocol": proto_score, "dvp": dvp_score, "dbdesign": dbd_score}
    best = max(scores, key=scores.get)
    if scores[best] > 0:
        role = best
        if proto_hits and best == "protocol":
            signals.append(f"content keywords: {proto_hits}")
        if dvp_hits and best == "dvp":
            signals.append(f"content keywords: {dvp_hits}")
        if dbd_hits and best == "dbdesign":
            signals.append(f"content keywords: {dbd_hits}")
        name_kw = name_hits(
            {"protocol": PROTOCOL_NAME_KW, "dvp": DVP_NAME_KW, "dbdesign": DBDESIGN_NAME_KW}[best]
        )
        if name_kw:
            signals.append(f"filename keywords: {name_kw}")

    return {
        "path": str(path),
        "filename": path.name,
        "format": fmt,
        "suggested_role": role,
        "scores": scores,
        "signals": signals,
    }


def _docx_sample_text(path: Path, max_paras: int = 200) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs[:max_paras]:
        parts.append(p.text)
    for tbl in doc.tables[:3]:
        for row in tbl.rows[:10]:
            parts.append(" ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _xlsx_sample_text(path: Path, max_rows: int = 30) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(ws.title)
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= max_rows:
                break
            parts.append(" ".join("" if v is None else str(v) for v in row))
    wb.close()
    return "\n".join(parts)


def _pdf_sample_text(path: Path, max_pages: int = 5) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages[:max_pages]:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def cmd_identify(args: argparse.Namespace) -> int:
    root = Path(args.directory)
    if not root.is_dir():
        fail(f"not a directory: {root}")
    records = []
    for path in sorted(root.rglob("*")):
        if path.suffix.lower() in (".docx", ".xlsx", ".pdf") and path.is_file():
            if path.name.startswith("~$"):  # office lock files
                continue
            records.append(classify_file(path))
    json.dump({"documents": records}, sys.stdout, ensure_ascii=False, indent=2)
    print()
    return 0


# ---------------------------------------------------------------------------
# extract
# ---------------------------------------------------------------------------

def load_doc_map(workdir: Path) -> dict:
    map_file = workdir / "doc-map.json"
    if map_file.exists():
        try:
            return json.loads(map_file.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            fail(f"corrupt doc-map.json: {map_file}")
    return {"documents": []}


def save_doc_map(workdir: Path, doc_map: dict) -> None:
    map_file = workdir / "doc-map.json"
    tmp = map_file.with_suffix(".json.tmp")
    tmp.write_text(json.dumps(doc_map, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(map_file)


def write_chunk(workdir: Path, doc_id: str, unit_id: str, locator: str, text: str) -> str:
    chunk_dir = workdir / "chunks" / doc_id
    chunk_dir.mkdir(parents=True, exist_ok=True)
    rel = f"chunks/{doc_id}/{unit_id}.md"
    header = f"<!-- doc_id: {doc_id} | unit_id: {unit_id} | locator: {locator} -->\n\n"
    (workdir / rel).write_text(header + text, encoding="utf-8")
    return rel


def split_long_text(paras: list[str], max_chars: int) -> list[str]:
    """Group paragraphs into blocks each <= max_chars (best effort)."""
    blocks: list[str] = []
    current: list[str] = []
    size = 0
    for para in paras:
        plen = len(para) + 1
        if current and size + plen > max_chars:
            blocks.append("\n\n".join(current))
            current, size = [], 0
        current.append(para)
        size += plen
    if current:
        blocks.append("\n\n".join(current))
    return blocks


def extract_docx(path: Path, doc_id: str, workdir: Path) -> list[dict]:
    from docx import Document

    doc = Document(str(path))
    # Walk paragraphs, tracking the heading stack. Build (heading_path, paras)
    # sections, then split oversized sections into paragraph-window chunks.
    sections: list[tuple[str, list[str]]] = []
    stack: list[str] = []  # heading titles by level order
    current_paras: list[str] = []
    current_path = "(front matter)"

    def flush() -> None:
        nonlocal current_paras
        paras = [p for p in current_paras if p.strip()]
        if paras:
            sections.append((current_path, paras))
        current_paras = []

    for para in doc.paragraphs:
        style = (para.style.name or "") if para.style else ""
        m = re.match(r"Heading\s*(\d+)", style, re.IGNORECASE)
        if m:
            level = int(m.group(1))
            flush()
            stack = stack[: level - 1]
            while len(stack) < level - 1:
                stack.append("")
            stack.append(para.text.strip())
            current_path = " > ".join(t for t in stack if t)
        else:
            current_paras.append(para.text)
    flush()

    # Include tables (docx tables are not part of paragraphs stream order in
    # python-docx; append them as their own section for completeness).
    if doc.tables:
        tbl_lines = []
        for ti, tbl in enumerate(doc.tables, start=1):
            tbl_lines.append(f"[Table {ti}]")
            for row in tbl.rows:
                tbl_lines.append(" | ".join(c.text.strip() for c in row.cells))
        sections.append(("(tables)", tbl_lines))

    units: list[dict] = []
    idx = 0
    for sec_path, paras in sections:
        blocks = split_long_text(paras, MAX_CHUNK_CHARS)
        for bi, block in enumerate(blocks):
            idx += 1
            unit_id = f"u{idx:03d}"
            locator = f"§{sec_path}"
            if len(blocks) > 1:
                locator = f"§{sec_path} [part {bi + 1}/{len(blocks)}]"
            chunk_file = write_chunk(workdir, doc_id, unit_id, locator, block)
            units.append({
                "unit_id": unit_id,
                "locator": locator,
                "chunk_file": chunk_file,
                "chars": len(block),
            })
    return units


def extract_xlsx(path: Path, doc_id: str, workdir: Path) -> tuple[list[dict], list[dict]]:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    units: list[dict] = []
    sheet_summaries: list[dict] = []
    idx = 0

    for ws in wb.worksheets:
        rows = list(ws.iter_rows(values_only=True))
        # Trim fully-empty trailing rows.
        while rows and all(v is None or str(v).strip() == "" for v in rows[-1]):
            rows.pop()
        n_rows = len(rows)
        n_cols = max((len(r) for r in rows), default=0)
        # Header detection: first non-empty row.
        header_row_idx = next(
            (i for i, r in enumerate(rows) if any(v is not None and str(v).strip() for v in r)),
            None,
        )
        headers: list[str] = []
        if header_row_idx is not None:
            headers = ["" if v is None else str(v).strip() for v in rows[header_row_idx]]
        sheet_summaries.append({
            "sheet": ws.title,
            "rows": n_rows,
            "cols": n_cols,
            "header_row": (header_row_idx + 1) if header_row_idx is not None else None,
            "headers": headers,
        })

        def row_text(r: tuple) -> str:
            return " | ".join("" if v is None else str(v) for v in r)

        if n_rows == 0:
            continue
        for start in range(0, n_rows, XLSX_ROWS_PER_CHUNK):
            end = min(start + XLSX_ROWS_PER_CHUNK, n_rows)
            idx += 1
            unit_id = f"u{idx:03d}"
            locator = f"<{ws.title}>!R{start + 1}-R{end}"
            lines = [f"# Sheet: {ws.title} rows {start + 1}-{end}"]
            if headers:
                lines.append("Headers: " + " | ".join(headers))
            for r in rows[start:end]:
                lines.append(row_text(r))
            text = "\n".join(lines)
            chunk_file = write_chunk(workdir, doc_id, unit_id, locator, text)
            units.append({
                "unit_id": unit_id,
                "locator": locator,
                "chunk_file": chunk_file,
                "chars": len(text),
            })
    wb.close()
    return units, sheet_summaries


def extract_pdf(path: Path, doc_id: str, workdir: Path) -> list[dict]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    n_pages = len(reader.pages)
    units: list[dict] = []
    idx = 0
    for start in range(0, n_pages, PDF_PAGES_PER_CHUNK):
        end = min(start + PDF_PAGES_PER_CHUNK, n_pages)
        idx += 1
        unit_id = f"u{idx:03d}"
        locator = f"p.{start + 1}-{end}"
        failed_pages = []
        parts = []
        for pi in range(start, end):
            try:
                txt = reader.pages[pi].extract_text() or ""
            except Exception:
                txt = ""
            if not txt.strip():
                failed_pages.append(pi + 1)
            parts.append(f"[p.{pi + 1}]\n{txt}")
        text = "\n\n".join(parts)
        chunk_file = write_chunk(workdir, doc_id, unit_id, locator, text)
        unit = {
            "unit_id": unit_id,
            "locator": locator,
            "chunk_file": chunk_file,
            "chars": len(text),
        }
        if failed_pages:
            unit["text_extraction_failed"] = True
            unit["failed_pages"] = failed_pages
        units.append(unit)
    return units


def cmd_extract(args: argparse.Namespace) -> int:
    path = Path(args.input)
    if not path.is_file():
        fail(f"input file not found: {path}")
    workdir = Path(args.workdir)
    workdir.mkdir(parents=True, exist_ok=True)
    doc_map = load_doc_map(workdir)

    suffix = path.suffix.lower()
    extra: dict = {}
    if suffix == ".docx":
        fmt = "docx"
        units = extract_docx(path, args.doc_id, workdir)
    elif suffix == ".xlsx":
        fmt = "xlsx"
        units, sheet_summaries = extract_xlsx(path, args.doc_id, workdir)
        extra["sheets"] = sheet_summaries
    elif suffix == ".pdf":
        fmt = "pdf"
        units = extract_pdf(path, args.doc_id, workdir)
    else:
        fail(f"unsupported format: {suffix} (expected .docx/.xlsx/.pdf)")

    # Idempotency: drop any existing entry for this doc_id and wipe its chunk
    # dir before/after writing (chunks already overwritten above; remove stale).
    doc_map["documents"] = [d for d in doc_map["documents"] if d.get("doc_id") != args.doc_id]
    entry = {
        "doc_id": args.doc_id,
        "role": args.role,
        "path": str(path),
        "format": fmt,
        "units": units,
    }
    entry.update(extra)
    doc_map["documents"].append(entry)
    save_doc_map(workdir, doc_map)

    # Remove stale chunk files for this doc that are no longer referenced.
    chunk_dir = workdir / "chunks" / args.doc_id
    if chunk_dir.is_dir():
        live = {u["chunk_file"] for u in units}
        for f in chunk_dir.glob("*.md"):
            if f"chunks/{args.doc_id}/{f.name}" not in live:
                f.unlink()

    failed = [u for u in units if u.get("text_extraction_failed")]
    json.dump({
        "doc_id": args.doc_id,
        "role": args.role,
        "format": fmt,
        "units": len(units),
        "text_extraction_failed_units": len(failed),
    }, sys.stdout, ensure_ascii=False, indent=2)
    print()
    if failed:
        eprint(f"warning: {len(failed)} unit(s) had pages with no extractable text; OCR/manual review may be needed")
    return 0


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Probe clinical trial documents: identify roles and extract text chunks."
    )
    sub = parser.add_subparsers(dest="command", required=True)

    p_id = sub.add_parser("identify", help="Scan a directory and classify documents.")
    p_id.add_argument("directory", help="Directory to scan (recursive).")
    p_id.set_defaults(func=cmd_identify)

    p_ex = sub.add_parser("extract", help="Extract one document into chunks + doc-map entry.")
    p_ex.add_argument("--input", required=True, help="Source file (.docx/.xlsx/.pdf).")
    p_ex.add_argument("--role", required=True, choices=["protocol", "dvp", "dbdesign"],
                      help="Document role (unknown roles must be confirmed by the caller first).")
    p_ex.add_argument("--doc-id", required=True, help="Stable document id (e.g. protocol, dvp).")
    p_ex.add_argument("--workdir", required=True, help="Working directory (e.g. .pd-extraction).")
    p_ex.set_defaults(func=cmd_extract)

    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    try:
        return args.func(args)
    except SystemExit:
        raise
    except Exception as exc:
        fail(f"{type(exc).__name__}: {exc}")
    return 1


if __name__ == "__main__":
    sys.exit(main())
